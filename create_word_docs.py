import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

dir_path = r"c:\code\Cecomtech\HuynhKiet"
os.makedirs(dir_path, exist_ok=True)

# Document 1: Theory
doc1 = Document()
heading = doc1.add_heading('LẬP TRÌNH PYTHON CHO NGƯỜI MỚI BẮT ĐẦU', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc1.add_paragraph('Mục tiêu: Làm quen với ngôn ngữ Python, hiểu cách máy tính hoạt động, biết cách cài đặt công cụ lập trình, viết chương trình đầu tiên và rèn luyện tư duy logic qua các lệnh rẽ nhánh.\n')

doc1.add_heading('Phần 1: Giới thiệu về Lập trình & Ngôn ngữ Python', level=1)
doc1.add_paragraph('Lập trình là gì? Về bản chất, lập trình là cách con người giao tiếp với máy tính. Chúng ta viết ra một tập hợp các chỉ dẫn (được gọi là mã nguồn - source code) để yêu cầu máy tính thực hiện các tác vụ cụ thể. Tuy nhiên, máy tính chỉ hiểu mã nhị phân (0 và 1), do đó chúng ta cần một "ngôn ngữ" trung gian và một "người phiên dịch" (Trình thông dịch - Interpreter) để dịch mã của chúng ta sang mã máy.')
p = doc1.add_paragraph('Python là ngôn ngữ lập trình bậc cao. Nó có cú pháp (syntax) vô cùng "dễ thương", ngắn gọn và cực kỳ gần gũi với ngôn ngữ tự nhiên (tiếng Anh). Thay vì viết hàng chục dòng code phức tạp như các ngôn ngữ khác, với Python, bạn đôi khi chỉ cần 1-2 dòng là đã xong việc.\n')
p.add_run('Python được ứng dụng ở đâu trong thực tế?\n').bold = True
doc1.add_paragraph('• Lập trình Web: Các hệ thống khổng lồ như Instagram, Spotify, Reddit, Netflix đều có sự đóng góp rất lớn của ngôn ngữ Python ở phía máy chủ (Backend).\n• Trí tuệ nhân tạo (AI) & Khoa học dữ liệu: ChatGPT, Midjourney, hệ thống gợi ý video của YouTube hay TikTok đều vận hành bằng Python.\n• Tự động hóa: Viết script tải hàng nghìn video, tự động trả lời email, lọc dữ liệu Excel trong chớp mắt.\n• Làm Game: Xây dựng các trò chơi 2D vui nhộn với thư viện PyGame.', style='List Bullet')

doc1.add_heading('Phần 2: Cài đặt và thiết lập Môi trường lập trình (VS Code)', level=1)
doc1.add_paragraph('Để viết mã Python, chúng ta cần 2 công cụ cốt lõi: (1) Bộ thông dịch Python (để máy tính hiểu code Python) và (2) Visual Studio Code (gọi tắt là VS Code - một công cụ soạn thảo xịn xò giúp tô màu code và gợi ý lệnh).')

doc1.add_heading('Bước 1: Cài đặt Python', level=2)
doc1.add_paragraph('- Truy cập trang chủ: python.org và tải xuống phiên bản mới nhất dành cho Windows.')
p = doc1.add_paragraph('- Chạy file vừa tải về. ')
p.add_run('LƯU Ý CỰC KỲ QUAN TRỌNG:').bold = True
p.add_run(' Ở màn hình cài đặt đầu tiên, bạn bắt buộc phải đánh dấu tích vào ô "Add Python to PATH" (hoặc "Add python.exe to PATH") ở dưới cùng. Sau đó bấm "Install Now" và đợi quá trình hoàn tất.')

doc1.add_heading('Bước 2: Cài đặt Visual Studio Code (VS Code)', level=2)
doc1.add_paragraph('- Truy cập trang chủ: code.visualstudio.com và tải VS Code về.')
doc1.add_paragraph('- Mở file cài đặt, nhấn "Next" (Tiếp tục) liên tục và chọn "Install". Việc cài đặt diễn ra như mọi phần mềm thông thường khác.')

doc1.add_heading('Bước 3: Cài đặt Extension Python cho VS Code', level=2)
doc1.add_paragraph('- Mở phần mềm VS Code vừa cài xong.')
doc1.add_paragraph('- Nhìn sang thanh dọc bên trái màn hình, nhấn vào biểu tượng 4 ô vuông (Extensions) hoặc dùng phím tắt Ctrl + Shift + X.')
doc1.add_paragraph('- Nhập chữ "Python" vào thanh tìm kiếm. Chọn Extension do chính "Microsoft" phát hành và bấm nút "Install".')

doc1.add_heading('Bước 4: Tạo và chạy chương trình đầu tiên trên VS Code', level=2)
doc1.add_paragraph('- Trên máy tính, tạo một thư mục mới có tên là "KhoaHoc_Python".')
doc1.add_paragraph('- Từ phần mềm VS Code, chọn File -> Open Folder (Mở thư mục) và chọn tới thư mục "KhoaHoc_Python" vừa tạo.')
doc1.add_paragraph('- Nhấn vào biểu tượng "New File" (Tạo file mới) hoặc nhấp đúp vào khoảng trống bên trái. Đặt tên file là "bai1.py" (Lưu ý: Bắt buộc phải có đuôi .py để VS Code hiểu đây là file Python).')
doc1.add_paragraph('- Gõ một dòng code ví dụ. Góc trên bên phải màn hình có một biểu tượng hình tam giác (nút Play/Run). Nhấp vào đó, chương trình sẽ chạy và kết quả sẽ hiển thị ở khu vực Terminal phía dưới.')

doc1.add_heading('Phần 3: Câu lệnh đầu tiên - Lên tiếng với thế giới (print)', level=1)
doc1.add_paragraph('Cách để bảo máy tính xuất/in một thông tin gì đó ra màn hình rất đơn giản, chúng ta dùng lệnh print() (có nghĩa là in ra).')
doc1.add_paragraph("print('Hello, World!')", style='Intense Quote')
doc1.add_paragraph('Chúng ta cũng có thể in nhiều thứ cùng lúc bằng cách dùng dấu phẩy:')
doc1.add_paragraph("print('Tên tôi là:', 'Hải', 'Năm nay', 18, 'tuổi')", style='Intense Quote')

doc1.add_heading('Phần 4: Biến số - Chiếc hộp ma thuật & Kiểu dữ liệu', level=1)
doc1.add_paragraph('Trong một hệ thống phần mềm, làm sao máy tính nhớ được bạn có bao nhiêu điểm, hay tên nhân vật là gì? Đó là nhờ "Biến" (Variable).')
doc1.add_paragraph('Biến giống như một chiếc hộp. Bạn dán nhãn cho hộp bằng một cái tên, và cất dữ liệu vào trong hộp đó để dùng lại sau này.')
doc1.add_paragraph("player_name = 'Yasuo'\nscore = 100\ncolor = 'Blue'\nlevel = 5\nis_alive = True", style='Intense Quote')
p = doc1.add_paragraph('Quy tắc đặt tên biến: Tên biến không được có dấu cách, không được bắt đầu bằng chữ số (1_name là sai), và chỉ nên dùng chữ cái tiếng Anh, số, dấu gạch dưới.\n')
p.add_run('Kiểu dữ liệu cơ bản:').bold = True
p.add_run('\n- Số nguyên (int): Các số không có phần thập phân như 100, 5, -20\n- Số thực (float): Số có phần thập phân như 3.14, -5.5\n- Chuỗi kí tự (str): Kí tự chữ, bắt buộc phải bọc trong nháy đơn \' \' hoặc nháy kép " ". Ví dụ: \'Yasuo\'\n- Logic (bool): Chỉ nhận 1 trong 2 giá trị là True (Đúng) hoặc False (Sai).')

doc1.add_heading('Phần 5: Phép toán cơ bản & "Phép biến hình" (Ép kiểu)', level=1)
doc1.add_paragraph('Python tính toán cực kỳ nhanh. Các phép toán cơ bản: Cộng (+), Trừ (-), Nhân (*), Chia (/), Chia lấy phần nguyên (//) và Chia lấy phần dư (%).')
doc1.add_paragraph("salary = 500\nbonus = 200\ntotal_money = salary + bonus\nprint(total_money)", style='Intense Quote')
doc1.add_paragraph('Phép biến hình (Ép kiểu - Type Casting): Đôi khi chúng ta có số bị bọc trong chuỗi \'50\', chúng ta không thể lấy chữ \'50\' cộng với số 20 được. Lúc này cần biến hình chuỗi thành số:')
doc1.add_paragraph("point = int('50') + 20  # Lệnh int() biến chuỗi '50' thành số 50. Kết quả là 70.", style='Intense Quote')

doc1.add_heading('Phần 6: Tương tác và nhập dữ liệu từ người dùng (input)', level=1)
doc1.add_paragraph('Thay vì gán cứng dữ liệu trong code, chúng ta có thể làm cho chương trình thông minh hơn bằng cách hỏi người dùng qua hàm input(). Khi gọi hàm này, chương trình sẽ tạm dừng và chờ người dùng gõ phím, ấn Enter để tiếp tục.')
doc1.add_paragraph("user_name = input('What is your name? ')\nprint('Nice to meet you, ' + user_name + '!')", style='Intense Quote')
doc1.add_paragraph('Lưu ý: Mọi thứ người dùng nhập vào từ hàm input() đều mặc định là Chuỗi (String). Nếu muốn dùng nó để tính toán, bạn phải ép kiểu nó sang số (ví dụ dùng int() hoặc float()).')

doc1.add_heading('Phần 7: Ra quyết định và rẽ nhánh (if - elif - else)', level=1)
doc1.add_paragraph('Trong đời sống, nếu trời mưa thì bạn sẽ ở nhà, nếu trời mát bạn sẽ đi dạo, còn lại bạn sẽ đi bơi. Máy tính cũng có thể rẽ nhánh dựa vào điều kiện nhờ cấu trúc if (nếu) - elif (nếu lại) - else (nếu không thì).')
doc1.add_paragraph("weather = 'Rain'\nif weather == 'Rain':\n    print('Stay at home!')\nelif weather == 'Cool':\n    print('Let us go for a walk!')\nelse:\n    print('Let us go swimming!')", style='Intense Quote')
doc1.add_paragraph('Lưu ý cú pháp quan trọng: Sau điều kiện phải có dấu hai chấm ":", và đoạn code bên dưới if/elif/else phải lùi lề vào trong (nhấn phím Tab). Điều này báo cho Python biết đoạn code đó thuộc về nhánh điều kiện phía trên.')
doc1.add_paragraph('Thực hành cuối buổi: Viết một chương trình hỏi điểm thi toán (math_score) của học sinh. Nếu điểm >= 8.0 in ra "Excellent", nếu điểm >= 5.0 in ra "Pass", ngược lại in ra "Fail".')

doc1.save(os.path.join(dir_path, 'Buoi_1_Ly_Thuyet_Python.docx'))

# Document 2: Multiple Choice
doc2 = Document()
heading2 = doc2.add_heading('BÀI TẬP TRẮC NGHIỆM PYTHON - BUỔI 1', 0)
heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc2.add_paragraph('Hãy chọn đáp án đúng nhất cho các câu hỏi sau đây:')

questions = [
    ("Câu 1: Câu lệnh nào sau đây dùng để in chữ 'Hello' ra màn hình?", ["A. print(Hello)", "B. print('Hello')", "C. output('Hello')", "D. echo 'Hello'"]),
    ("Câu 2: Python KHÔNG thể làm được việc gì trong các việc sau?", ["A. Tạo trang web backend", "B. Lập trình trí tuệ nhân tạo (AI)", "C. Nấu một bữa cơm thực tế", "D. Tự động hóa tác vụ Excel"]),
    ("Câu 3: Đâu là cách tạo biến (chiếc hộp) đúng để lưu tuổi của người chơi?", ["A. 18 = player_age", "B. player_age = 18", "C. player_age == 18", "D. player_age : 18"]),
    ("Câu 4: Kết quả của đoạn code sau là gì?\nprint(5 + 3 * 2)", ["A. 16", "B. 11", "C. 10", "D. Lỗi"]),
    ("Câu 5: Khi lưu tên nhân vật (ví dụ: 'Batman'), bạn phải đặt nó trong dấu gì?", ["A. Dấu ngoặc vuông [ ]", "B. Dấu ngoặc nhọn { }", "C. Dấu ngoặc đơn ' ' hoặc ngoặc kép \" \"", "D. Không cần dấu gì cả"]),
    ("Câu 6: Hàm nào được sử dụng để hỏi người dùng nhập thông tin từ bàn phím?", ["A. get()", "B. scan()", "C. keyboard()", "D. input()"]),
    ("Câu 7: Kết quả của câu lệnh sau là gì?\nprint('Ha' + 'Ha')", ["A. Ha Ha", "B. HaHa", "C. 2Ha", "D. Lỗi"]),
    ("Câu 8: IDE là phần mềm dùng để làm gì?", ["A. Chơi game 3D", "B. Nghe nhạc", "C. Giúp lập trình viên viết và chạy code", "D. Diệt virus"]),
    ("Câu 9: Trong Python, ký hiệu nào dùng để chia lấy phần nguyên?", ["A. /", "B. %", "C. //", "D. **"]),
    ("Câu 10: Đoạn code sau bị lỗi ở đâu?\nhero_name = 'Superman\nprint(hero_name)", ["A. Thiếu dấu nháy đóng lại ở chữ 'Superman'", "B. print viết sai", "C. Tên biến không hợp lệ", "D. Không có lỗi"]),
    ("Câu 11: Kiểu dữ liệu số nguyên trong Python được gọi là gì?", ["A. string", "B. boolean", "C. int", "D. float"]),
    ("Câu 12: Đâu là tên biến hợp lệ trong Python?", ["A. 1st_player", "B. player name", "C. player_name", "D. player-name"]),
    ("Câu 13: Đoạn code sau sẽ in ra màn hình nội dung gì?\nscore = 10\nscore = score + 5\nprint(score)", ["A. 10", "B. 5", "C. 15", "D. score + 5"]),
    ("Câu 14: Lệnh nào sau đây chuyển đổi chuỗi '100' thành số nguyên 100?", ["A. to_int('100')", "B. float('100')", "C. number('100')", "D. int('100')"]),
    ("Câu 15: Cho biến is_game_over = False. Biến này thuộc kiểu dữ liệu nào?", ["A. int", "B. str", "C. bool", "D. float"]),
    ("Câu 16: Nếu muốn tính 2 mũ 3 trong Python, ta dùng toán tử nào?", ["A. 2 ^ 3", "B. 2 ** 3", "C. 2 * 3", "D. pow 2 3"]),
    ("Câu 17: Kết quả của đoạn code sau là gì?\nx = 10\nif x > 5:\n    print('Big')\nelse:\n    print('Small')", ["A. Small", "B. Lỗi", "C. Big", "D. Không in ra gì"]),
    ("Câu 18: Lệnh input() luôn trả về kiểu dữ liệu gì cho dù người dùng có nhập số đi chăng nữa?", ["A. int", "B. float", "C. bool", "D. str (Chuỗi)"]),
    ("Câu 19: Dấu `#` trong Python được dùng để làm gì?", ["A. Bắt đầu một câu lệnh", "B. Viết chú thích (comment) mà máy tính sẽ bỏ qua", "C. Kết thúc chương trình", "D. Tính toán logic"]),
    ("Câu 20: Bạn viết chương trình yêu cầu người dùng nhập số điểm. Lệnh nào sau đây là cách kết hợp đúng giữa nhập và ép kiểu?", ["A. score = int(input('Enter score: '))", "B. score = input(int('Enter score: '))", "C. score = input('Enter score: ').int()", "D. int score = input('Enter score: ')"])
]

for q, opts in questions:
    doc2.add_paragraph(q, style='List Number')
    for opt in opts:
        doc2.add_paragraph(opt)
    doc2.add_paragraph()

doc2.add_page_break()
doc2.add_heading('ĐÁP ÁN', level=1)
answers = [
    "Câu 1: B", "Câu 2: C", "Câu 3: B", "Câu 4: B (Nhân chia trước, cộng trừ sau)", 
    "Câu 5: C", "Câu 6: D", "Câu 7: B", "Câu 8: C", "Câu 9: C", "Câu 10: A",
    "Câu 11: C", "Câu 12: C", "Câu 13: C", "Câu 14: D", "Câu 15: C", "Câu 16: B",
    "Câu 17: C", "Câu 18: D", "Câu 19: B", "Câu 20: A"
]

for ans in answers:
    doc2.add_paragraph(ans)

doc2.save(os.path.join(dir_path, 'Buoi_1_Trac_Nghiem_Python.docx'))
print("Files updated successfully.")
