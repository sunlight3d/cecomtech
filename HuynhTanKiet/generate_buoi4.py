import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(20, 100, 20)  # Dark green
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p.paragraph_format.left_indent = Pt(20)

doc_lt = docx.Document()

title = doc_lt.add_heading('LẬP TRÌNH PYTHON: BUỔI 4 - XỬ LÝ LỖI, LÀM VIỆC VỚI FILE VÀ OOP', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

p_muc_tieu = doc_lt.add_paragraph()
p_muc_tieu.add_run('Mục tiêu: ').bold = True
p_muc_tieu.add_run('Biết cách bắt lỗi và xử lý ngoại lệ (Try/Except) để chương trình không bị crash. Thành thạo việc đọc và ghi dữ liệu ra File. Bước đầu làm quen với Lập trình Hướng đối tượng (OOP) qua Khái niệm Lớp (Class) và Đối tượng (Object). (Thời lượng: 6 giờ)\n')

# Phần 1
h1 = doc_lt.add_heading('Phần 1: Xử lý ngoại lệ (Exception Handling)', level=2)
for run in h1.runs: run.font.name = 'Arial'

doc_lt.add_paragraph('1.1 Lỗi (Bug) và Ngoại lệ (Exception) là gì?')
doc_lt.add_paragraph('Khi chương trình đang chạy mà gặp phải một phép tính vô lý (chia cho 0) hoặc không tìm thấy file, chương trình sẽ báo lỗi và dừng lại ngay lập tức (crash). Xử lý ngoại lệ giúp chương trình báo lỗi một cách "thân thiện" và tiếp tục chạy.')

doc_lt.add_paragraph('1.2 Cú pháp Try...Except')
doc_lt.add_paragraph('Đặt đoạn code có nguy cơ bị lỗi vào khối `try`. Nếu có lỗi xảy ra, khối `except` sẽ bắt lấy lỗi đó thay vì làm sập chương trình.')
add_code_block(doc_lt, "try:\n    result = 10 / 0\nexcept ZeroDivisionError:\n    print('Lỗi: Không thể chia cho 0!')")

doc_lt.add_paragraph('1.3 Khối Finally')
doc_lt.add_paragraph('Khối `finally` luôn luôn được chạy bất kể có lỗi hay không. Thường dùng để dọn dẹp bộ nhớ hoặc đóng file.')
add_code_block(doc_lt, "try:\n    print('Đang xử lý...')\nfinally:\n    print('Luôn chạy câu lệnh này')")

# Phần 2
h2 = doc_lt.add_heading('Phần 2: Làm việc với File (File Handling)', level=2)
for run in h2.runs: run.font.name = 'Arial'

doc_lt.add_paragraph('2.1 Mở và Đọc File')
doc_lt.add_paragraph('Chúng ta dùng hàm `open()` với chế độ "r" (read) để đọc file. Tốt nhất là dùng từ khóa `with` để tự động đóng file sau khi đọc xong.')
add_code_block(doc_lt, "with open('data.txt', 'r') as file:\n    content = file.read()\n    print(content)")

doc_lt.add_paragraph('2.2 Ghi dữ liệu vào File')
doc_lt.add_paragraph('Dùng chế độ "w" (write) để ghi đè file mới, hoặc "a" (append) để viết tiếp vào cuối file đã có.')
add_code_block(doc_lt, "with open('log.txt', 'a') as file:\n    file.write('Dòng log mới\\n')")

# Phần 3
h3 = doc_lt.add_heading('Phần 3: Lập trình Hướng đối tượng cơ bản (OOP)', level=2)
for run in h3.runs: run.font.name = 'Arial'

doc_lt.add_paragraph('3.1 Khái niệm Class và Object')
doc_lt.add_paragraph('Class (Lớp) là một bản thiết kế (ví dụ: bản vẽ ô tô). Object (Đối tượng) là một thực thể được tạo ra từ bản thiết kế đó (ví dụ: chiếc ô tô Toyota ngoài đời thực).')

doc_lt.add_paragraph('3.2 Khởi tạo Class và phương thức (Method)')
doc_lt.add_paragraph('Dùng từ khóa `class`. Phương thức đặc biệt `__init__` (gọi là constructor) được tự động chạy khi ta tạo một Object mới. Từ khóa `self` đại diện cho chính đối tượng đó.')
add_code_block(doc_lt, "class Car:\n    def __init__(self, brand, color):\n        self.brand = brand\n        self.color = color\n\n    def drive(self):\n        print(f'Chiếc xe {self.color} của {self.brand} đang chạy!')\n\n# Tạo Object\nmy_car = Car('Toyota', 'Đen')\nmy_car.drive()")

doc_lt.save('Buoi_4_Ly_Thuyet_Python.docx')

# --- Tạo file Trắc Nghiệm ---
doc_tn = docx.Document()

title_tn = doc_tn.add_heading('BÀI TẬP TRẮC NGHIỆM PYTHON - BUỔI 4 (50 CÂU)', level=1)
title_tn.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title_tn.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

doc_tn.add_paragraph('Hãy chọn đáp án đúng nhất cho các câu hỏi sau đây:\n')

cau_hoi = [
    # Try/Except (1-15)
    {"q": "Câu 1: Khối lệnh nào dùng để bắt và xử lý ngoại lệ trong Python?", "options": ["A. catch / throw", "B. try / except", "C. try / catch", "D. error / handle"]},
    {"q": "Câu 2: Đoạn mã có khả năng gây ra lỗi nên được đặt trong khối lệnh nào?", "options": ["A. except", "B. try", "C. finally", "D. else"]},
    {"q": "Câu 3: Khối lệnh `finally` có đặc điểm gì?", "options": ["A. Chỉ chạy khi có lỗi", "B. Chỉ chạy khi không có lỗi", "C. Luôn luôn chạy dù có lỗi hay không", "D. Dùng để sinh ra lỗi nhân tạo"]},
    {"q": "Câu 4: Tên của lỗi xảy ra khi bạn cố gắng chia một số cho 0 là gì?", "options": ["A. ValueError", "B. TypeError", "C. ZeroDivisionError", "D. DivideByZero"]},
    {"q": "Câu 5: Từ khóa nào dùng để CỐ TÌNH tạo ra (raise) một ngoại lệ trong Python?", "options": ["A. throw", "B. error", "C. exception", "D. raise"]},
    {"q": "Câu 6: Nếu không dùng `except` với một tên lỗi cụ thể (ví dụ: `except:`), nó sẽ:", "options": ["A. Chỉ bắt các lỗi cú pháp", "B. Bắt tất cả các loại lỗi", "C. Không bắt được lỗi nào", "D. Gây ra lỗi cú pháp"]},
    {"q": "Câu 7: Lỗi `ValueError` thường xảy ra khi nào?", "options": ["A. Biến chưa được khai báo", "B. Chia cho 0", "C. Hàm nhận đúng kiểu dữ liệu nhưng sai giá trị (vd: int('abc'))", "D. Cú pháp code sai"]},
    {"q": "Câu 8: Lỗi `TypeError` thường xảy ra khi nào?", "options": ["A. Khi tính tổng chuỗi và số: 'hello' + 5", "B. Khi gọi sai tên hàm", "C. Khi mở file không tồn tại", "D. Khi bộ nhớ bị đầy"]},
    {"q": "Câu 9: `except Exception as e:` dùng để làm gì?", "options": ["A. In ra chữ 'e'", "B. Gán thông tin lỗi vào biến 'e' để dễ dàng in ra", "C. Bỏ qua tất cả các lỗi", "D. Tạo ra một ngoại lệ mới tên là e"]},
    {"q": "Câu 10: Có thể sử dụng nhiều khối `except` cho một khối `try` duy nhất không?", "options": ["A. Có", "B. Không"]},
    {"q": "Câu 11: Lỗi `FileNotFoundError` xảy ra khi:", "options": ["A. Ổ cứng bị hỏng", "B. Tên file không hợp lệ", "C. File cố gắng mở không tồn tại trên máy", "D. File không có đuôi .txt"]},
    {"q": "Câu 12: Khối `else` khi đi kèm với `try / except` sẽ chạy khi nào?", "options": ["A. Khi CÓ lỗi xảy ra", "B. Khi KHÔNG CÓ lỗi nào xảy ra trong khối try", "C. Luôn luôn chạy", "D. Không tồn tại khối else trong try/except"]},
    {"q": "Câu 13: Đâu là một ngoại lệ hợp lệ trong Python?", "options": ["A. IndexOutOfBounds", "B. IndexError", "C. ArrayError", "D. OutOfIndex"]},
    {"q": "Câu 14: Lỗi `KeyError` thường xuất hiện khi thao tác với cấu trúc dữ liệu nào?", "options": ["A. List", "B. Tuple", "C. Dictionary", "D. Set"]},
    {"q": "Câu 15: Nếu muốn bỏ qua hoàn toàn một lỗi (không in gì ra), ta dùng câu lệnh gì trong `except`?", "options": ["A. continue", "B. ignore", "C. skip", "D. pass"]},

    # File Handling (16-30)
    {"q": "Câu 16: Hàm nào được sử dụng để mở một file trong Python?", "options": ["A. open()", "B. read()", "C. open_file()", "D. load()"]},
    {"q": "Câu 17: Chế độ (mode) nào dùng để ĐỌC file?", "options": ["A. 'w'", "B. 'a'", "C. 'r'", "D. 'x'"]},
    {"q": "Câu 18: Chế độ (mode) nào dùng để GHI file và sẽ GHI ĐÈ toàn bộ nội dung cũ?", "options": ["A. 'r'", "B. 'w'", "C. 'a'", "D. 'r+'"]},
    {"q": "Câu 19: Chế độ 'a' (append) dùng để làm gì?", "options": ["A. Đọc file từ đầu", "B. Ghi đè file", "C. Ghi thêm dữ liệu vào cuối file mà không làm mất dữ liệu cũ", "D. Xóa file"]},
    {"q": "Câu 20: Tại sao chúng ta nên luôn gọi `.close()` sau khi thao tác với file xong (nếu không dùng `with`)?", "options": ["A. Để giảm dung lượng file", "B. Để giải phóng tài nguyên hệ thống và đảm bảo dữ liệu được ghi xuống đĩa", "C. Để mã hóa file", "D. Để tránh bị virus"]},
    {"q": "Câu 21: Câu lệnh `with open('data.txt', 'r') as file:` có lợi ích lớn nhất là gì?", "options": ["A. Làm code chạy nhanh gấp đôi", "B. Tự động đóng file sau khi ra khỏi khối lệnh (kể cả khi có lỗi)", "C. Tự động tạo file nếu file chưa tồn tại", "D. Mã hóa nội dung file"]},
    {"q": "Câu 22: Phương thức `.read()` làm nhiệm vụ gì?", "options": ["A. Đọc file theo từng dòng", "B. Đọc toàn bộ nội dung file dưới dạng một chuỗi (String) duy nhất", "C. Đọc ký tự đầu tiên", "D. Đọc từng từ một"]},
    {"q": "Câu 23: Phương thức `.readlines()` trả về kiểu dữ liệu gì?", "options": ["A. String", "B. Tuple", "C. List (mỗi dòng là một phần tử)", "D. Dictionary"]},
    {"q": "Câu 24: Phương thức nào dùng để ĐỌC MỘT DÒNG từ file?", "options": ["A. read()", "B. readline()", "C. readlines()", "D. read_one()"]},
    {"q": "Câu 25: Phương thức nào dùng để ghi chuỗi dữ liệu vào file?", "options": ["A. write()", "B. writeline()", "C. append()", "D. put()"]},
    {"q": "Câu 26: Nếu mở một file bằng chế độ 'w' nhưng file đó CHƯA TỒN TẠI, Python sẽ làm gì?", "options": ["A. Báo lỗi FileNotFoundError", "B. Tự động tạo một file mới với tên đó", "C. Bỏ qua câu lệnh", "D. Chờ người dùng tạo file"]},
    {"q": "Câu 27: Nếu mở một file bằng chế độ 'r' nhưng file đó CHƯA TỒN TẠI, Python sẽ làm gì?", "options": ["A. Báo lỗi FileNotFoundError", "B. Tự động tạo file mới", "C. Trả về giá trị None", "D. Đợi đến khi có file"]},
    {"q": "Câu 28: Ký tự nào thường dùng để ngắt dòng (xuống dòng) khi ghi dữ liệu text trong Python?", "options": ["A. \\t", "B. \\r", "C. \\n", "D. \\b"]},
    {"q": "Câu 29: Để đọc và ghi các định dạng dữ liệu phức tạp như List, Dictionary ra file chuẩn, ta dùng thư viện nào?", "options": ["A. sys", "B. json", "C. math", "D. csv"]},
    {"q": "Câu 30: Chữ 'b' trong chế độ mở file (ví dụ 'rb' hoặc 'wb') có nghĩa là gì?", "options": ["A. bold", "B. backup", "C. block", "D. binary (nhị phân, dùng cho hình ảnh/video)"]},

    # OOP (31-40)
    {"q": "Câu 31: Khái niệm 'Class' trong OOP là gì?", "options": ["A. Một biến lưu trữ dữ liệu", "B. Một hàm thực thi lệnh", "C. Một khuôn mẫu (bản thiết kế) để tạo ra các đối tượng", "D. Một thư viện bên ngoài"]},
    {"q": "Câu 32: 'Object' (Đối tượng) là gì?", "options": ["A. Là một thể hiện (instance) cụ thể được tạo ra từ Class", "B. Là tên gọi khác của Class", "C. Là một hàm ẩn danh", "D. Là hệ điều hành"]},
    {"q": "Câu 33: Phương thức (Method) đặc biệt nào được dùng làm Hàm khởi tạo (Constructor) trong Python?", "options": ["A. __constructor__", "B. __init__", "C. __start__", "D. __main__"]},
    {"q": "Câu 34: Tham số ĐẦU TIÊN của mọi phương thức bên trong một Class thông thường là gì?", "options": ["A. this", "B. object", "C. self", "D. master"]},
    {"q": "Câu 35: Từ khóa `self` trong OOP Python có ý nghĩa gì?", "options": ["A. Khai báo biến toàn cục", "B. Đại diện cho chính đối tượng đang gọi phương thức đó", "C. Đại diện cho Class cha", "D. Một thư viện bắt buộc"]},
    {"q": "Câu 36: Tính chất nào KHÔNG phải là một trong 4 tính chất cơ bản của Lập trình hướng đối tượng?", "options": ["A. Tính Kế thừa (Inheritance)", "B. Tính Đa hình (Polymorphism)", "C. Tính Toàn cục (Globalism)", "D. Tính Đóng gói (Encapsulation)"]},
    {"q": "Câu 37: Để Class `Dog` kế thừa từ Class `Animal`, cú pháp nào đúng?", "options": ["A. class Dog inherits Animal:", "B. class Dog(Animal):", "C. class Dog extends Animal:", "D. class Animal(Dog):"]},
    {"q": "Câu 38: Phương thức (Method) thực chất là gì?", "options": ["A. Là một biến nằm trong Class", "B. Là một vòng lặp", "C. Là một hàm (Function) được định nghĩa bên trong một Class", "D. Là một Dictionary"]},
    {"q": "Câu 39: Để tạo một object tên là `my_cat` từ class `Cat`, ta viết:", "options": ["A. my_cat = new Cat()", "B. Cat my_cat = Cat()", "C. my_cat = Cat()", "D. create Cat my_cat"]},
    {"q": "Câu 40: Nếu không viết hàm `__init__` cho một Class, Python sẽ phản ứng thế nào?", "options": ["A. Báo lỗi cú pháp", "B. Class đó không thể tạo object", "C. Cung cấp một hàm `__init__` mặc định (rỗng) ngầm định", "D. Yêu cầu nhập từ bàn phím"]},

    # Đoán kết quả Code (41-50)
    {"q": "Câu 41: Kết quả đầu ra của đoạn code sau là gì?", "code": "try:\n    print(10 / 2)\nexcept ZeroDivisionError:\n    print('Lỗi chia 0')\nfinally:\n    print('Xong!')", "options": ["A. 5.0\\nXong!", "B. Lỗi chia 0\\nXong!", "C. 5.0", "D. Xong!"]},
    {"q": "Câu 42: Kết quả đầu ra của đoạn code sau là gì?", "code": "try:\n    int('Hello')\nexcept ValueError:\n    print('Sai giá trị')\nexcept TypeError:\n    print('Sai kiểu')", "options": ["A. Sai kiểu", "B. Sai giá trị", "C. Không in gì", "D. Báo lỗi crash"]},
    {"q": "Câu 43: Kết quả đầu ra của đoạn code sau là gì?", "code": "class Person:\n    def __init__(self, name):\n        self.name = name\n\np = Person('John')\nprint(p.name)", "options": ["A. name", "B. p.name", "C. John", "D. Lỗi"]},
    {"q": "Câu 44: Kết quả đầu ra của đoạn code sau là gì?", "code": "def divide(a, b):\n    try:\n        return a / b\n    except:\n        return 'Lỗi'\n\nprint(divide(10, 0))", "options": ["A. Lỗi", "B. 0", "C. Báo lỗi crash (Exception)", "D. None"]},
    {"q": "Câu 45: Kết quả đầu ra của đoạn code sau là gì?", "code": "my_dict = {'a': 1, 'b': 2}\ntry:\n    print(my_dict['c'])\nexcept KeyError:\n    print('Khóa không tồn tại')", "options": ["A. None", "B. 0", "C. Báo lỗi crash", "D. Khóa không tồn tại"]},
    {"q": "Câu 46: Kết quả đầu ra của đoạn code sau là gì?", "code": "x = [1, 2, 3]\ntry:\n    x[5] = 10\nexcept IndexError:\n    print('Quá giới hạn')\nexcept Exception:\n    print('Lỗi khác')", "options": ["A. 10", "B. Quá giới hạn", "C. Lỗi khác", "D. Code chạy không lỗi nhưng không thay đổi x"]},
    {"q": "Câu 47: Kết quả đầu ra của đoạn code sau là gì?", "code": "with open('temp.txt', 'w') as f:\n    f.write('Python')\n\nwith open('temp.txt', 'r') as f:\n    print(f.read())", "options": ["A. Python", "B. FileNotFoundError", "C. Báo lỗi do quên close file", "D. Trống"]},
    {"q": "Câu 48: Kết quả đầu ra của đoạn code sau là gì?", "code": "class Calculator:\n    def add(self, a, b):\n        return a + b\n\ncalc = Calculator()\nprint(calc.add(3, 4))", "options": ["A. 7", "B. 34", "C. None", "D. Lỗi thiếu self khi gọi"]},
    {"q": "Câu 49: Kết quả đầu ra của đoạn code sau là gì?", "code": "num = '5'\ntry:\n    result = num + 5\n    print(result)\nexcept TypeError:\n    print('Type Error')\nexcept ValueError:\n    print('Value Error')", "options": ["A. 10", "B. 55", "C. Type Error", "D. Value Error"]},
    {"q": "Câu 50: Kết quả đầu ra của đoạn code sau là gì?", "code": "def process():\n    try:\n        return 1\n    finally:\n        return 2\n\nprint(process())", "options": ["A. 1", "B. 2", "C. Trả về cả 1 và 2", "D. Lỗi cú pháp"]}
]

for ch in cau_hoi:
    p = doc_tn.add_paragraph(ch['q'])
    p.runs[0].bold = True
    
    if "code" in ch:
        add_code_block(doc_tn, ch["code"])
        
    for opt in ch['options']:
        doc_tn.add_paragraph(opt)
    doc_tn.add_paragraph() # Dòng trống

doc_tn.save('Buoi_4_Trac_Nghiem_Python.docx')
