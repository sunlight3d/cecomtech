import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

dir_path = r"c:\code\Cecomtech\HuynhKiet"
os.makedirs(dir_path, exist_ok=True)

# Forest Canopy Theme
# Colors
HEADING_COLOR = RGBColor(0x2E, 0x4E, 0x3F) # Deep Forest Green
BODY_COLOR = RGBColor(0x3E, 0x3E, 0x3E)    # Dark Gray/Brown
CODE_COLOR = RGBColor(0x55, 0x6B, 0x2F)    # Dark Olive Green

def apply_heading_style(heading, level):
    for run in heading.runs:
        run.font.name = 'Georgia'
        run.font.color.rgb = HEADING_COLOR
        if level == 0:
            run.font.size = Pt(20)
            run.bold = True
        elif level == 1:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True

def add_body_paragraph(doc, text, bold_prefix=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Arial'
        r.font.color.rgb = BODY_COLOR
        
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.color.rgb = BODY_COLOR
    else:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.color.rgb = BODY_COLOR
    return p

def add_code_paragraph(doc, text):
    p = doc.add_paragraph(style='Intense Quote')
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.color.rgb = CODE_COLOR
    r.bold = True
    return p

# --- DOCUMENT 1: THEORY ---
doc1 = Document()
h0 = doc1.add_heading('LẬP TRÌNH PYTHON CHO NGƯỜI MỚI BẮT ĐẦU - BUỔI 2', 0)
h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
apply_heading_style(h0, 0)

add_body_paragraph(doc1, 'Mục tiêu: Nắm vững cấu trúc dữ liệu Danh sách (List), sử dụng thành thạo Vòng lặp (For, While) để tối ưu hóa mã nguồn, giải quyết các bài toán lặp đi lặp lại. (Thời lượng: 4 giờ)\n')

h1 = doc1.add_heading('Phần 1: Danh sách (List) - Chiếc túi thần kỳ (60 phút)', level=1)
apply_heading_style(h1, 1)
add_body_paragraph(doc1, 'Trong buổi 1, chúng ta dùng "Biến" như một chiếc hộp chỉ chứa được 1 món đồ. Vậy nếu bạn muốn quản lý 100 học sinh, bạn phải tạo 100 biến? Không! Chúng ta sẽ dùng "Danh sách" (List) - giống như một chiếc ba lô chứa được vô số món đồ.')
add_code_paragraph(doc1, "inventory = ['Kiếm', 'Khiên', 'Bình máu', 'Bản đồ']")
add_body_paragraph(doc1, 'Cách sử dụng List:', '1. Khởi tạo: ')
add_body_paragraph(doc1, 'Mỗi món đồ trong ba lô được đánh số thứ tự từ 0. Món đầu tiên là số 0, món thứ hai là số 1... Số thứ tự này gọi là Chỉ số (Index).', '2. Lấy đồ vật (Truy cập Index): ')
add_code_paragraph(doc1, "print(inventory[0])  # In ra 'Kiếm'\nprint(inventory[2])  # In ra 'Bình máu'")
add_body_paragraph(doc1, 'Dùng hàm append() để nhét thêm đồ vào cuối ba lô.', '3. Thêm đồ vật (Append): ')
add_code_paragraph(doc1, "inventory.append('Chìa khóa')")
add_body_paragraph(doc1, 'Dùng lệnh remove() để bỏ một món đồ ra khỏi ba lô, hoặc pop() để vứt món đồ cuối cùng.', '4. Xóa đồ vật: ')

h1 = doc1.add_heading('Phần 2: Vòng lặp FOR - Cỗ máy tự động hóa (60 phút)', level=1)
apply_heading_style(h1, 1)
add_body_paragraph(doc1, 'Khi bạn muốn máy tính làm một việc gì đó nhiều lần, thay vì viết lệnh copy-paste hàng trăm lần, ta dùng vòng lặp.')
add_body_paragraph(doc1, 'Vòng lặp For đặc biệt hữu ích khi duyệt qua từng món đồ trong một Danh sách (List).')
add_code_paragraph(doc1, "for item in inventory:\n    print('Tôi đang mang theo:', item)")
add_body_paragraph(doc1, 'Lệnh range(n) tạo ra các số từ 0 đến n-1. Nếu bạn muốn in chữ "Hello" 5 lần:', 'Hàm range(): ')
add_code_paragraph(doc1, "for i in range(5):\n    print('Hello')")
add_body_paragraph(doc1, 'Thực hành: Viết chương trình tạo một List chứa tên 3 người bạn. Dùng vòng lặp for in ra câu chúc: "Happy New Year, [tên bạn]!".')

h1 = doc1.add_heading('Phần 3: Vòng lặp WHILE - Làm cho đến khi dừng (60 phút)', level=1)
apply_heading_style(h1, 1)
add_body_paragraph(doc1, 'Vòng lặp While hoạt động dựa trên điều kiện: Chừng nào điều kiện còn Đúng (True), thì tiếp tục làm. Nó giống như việc "Chừng nào bụng còn đói, thì tiếp tục ăn".')
add_code_paragraph(doc1, "energy = 0\nwhile energy < 100:\n    print('Đang sạc pin...', energy, '%')\n    energy = energy + 20\nprint('Pin đã đầy!')")
add_body_paragraph(doc1, 'Trong vòng lặp, bạn có thể dùng lệnh break để "đập vỡ" vòng lặp ngay lập tức, và continue để bỏ qua vòng hiện tại, nhảy sang vòng tiếp theo.', 'Lệnh break và continue: ')
add_code_paragraph(doc1, "while True:  # Vòng lặp vô tận\n    answer = input('Nhập mật khẩu: ')\n    if answer == '1234':\n        print('Mở khóa!')\n        break  # Thoát khỏi vòng lặp")

h1 = doc1.add_heading('Phần 4: Dự án Mini - Cửa hàng thần bí (60 phút)', level=1)
apply_heading_style(h1, 1)
add_body_paragraph(doc1, 'Yêu cầu: Kết hợp List, While và If-Else để viết chương trình:')
add_body_paragraph(doc1, '- Máy tính cho một danh sách sản phẩm: products = ["Áo", "Quần", "Mũ"]')
add_body_paragraph(doc1, '- Sử dụng vòng lặp while True để liên tục hỏi người dùng: "Bạn muốn mua gì? (gõ chữ q để thoát)"')
add_body_paragraph(doc1, '- Nếu người dùng gõ món đồ có trong danh sách, in ra "Đã mua thành công!" và xóa món đó khỏi danh sách.')
add_body_paragraph(doc1, '- Nếu người dùng gõ món không có, in ra "Hết hàng!".')
add_body_paragraph(doc1, '- Nếu gõ "q", dùng lệnh break để dừng chương trình.')

doc1.save(os.path.join(dir_path, 'Buoi_2_Ly_Thuyet_Python.docx'))

# --- DOCUMENT 2: MULTIPLE CHOICE ---
doc2 = Document()
h0 = doc2.add_heading('BÀI TẬP TRẮC NGHIỆM PYTHON - BUỔI 2', 0)
h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
apply_heading_style(h0, 0)

add_body_paragraph(doc2, 'Hãy chọn đáp án đúng nhất cho các câu hỏi sau đây:')

questions = [
    ("Câu 1: List (Danh sách) trong Python được khai báo bằng cặp dấu ngoặc nào?", ["A. ( )", "B. [ ]", "C. { }", "D. < >"]),
    ("Câu 2: Khai báo List nào sau đây là hợp lệ?", ["A. numbers = [1, 2, 3]", "B. numbers = (1, 2, 3)", "C. numbers = 1, 2, 3", "D. numbers = {1, 2, 3}"]),
    ("Câu 3: Trong List `fruits = ['Táo', 'Cam', 'Nho']`, lệnh `print(fruits[1])` sẽ in ra gì?", ["A. Táo", "B. Nho", "C. Cam", "D. Lỗi"]),
    ("Câu 4: Chỉ số (Index) của phần tử đầu tiên trong List là số mấy?", ["A. 1", "B. 0", "C. -1", "D. 2"]),
    ("Câu 5: Làm thế nào để thêm 'Chuối' vào cuối List `fruits`?", ["A. fruits.add('Chuối')", "B. fruits.append('Chuối')", "C. fruits.insert('Chuối')", "D. fruits.push('Chuối')"]),
    ("Câu 6: Lệnh `remove('Táo')` dùng để làm gì?", ["A. Xóa tất cả các chữ Táo trong máy tính", "B. Xóa 'Táo' khỏi List", "C. Tìm vị trí của chữ 'Táo'", "D. Thêm chữ 'Táo' vào List"]),
    ("Câu 7: Vòng lặp `for` thường được sử dụng khi nào?", ["A. Khi biết trước số lần lặp hoặc duyệt qua một danh sách.", "B. Khi không biết trước số lần lặp.", "C. Khi muốn máy tính bị treo.", "D. Khi muốn tính tổng."]),
    ("Câu 8: Hàm `range(5)` sẽ tạo ra các số nào?", ["A. 1, 2, 3, 4, 5", "B. 0, 1, 2, 3, 4", "C. 0, 1, 2, 3, 4, 5", "D. 1, 2, 3, 4"]),
    ("Câu 9: Kết quả của đoạn code sau là gì?\nfor i in range(2):\n    print('Hello')", ["A. In chữ Hello 1 lần", "B. In chữ Hello 2 lần", "C. In chữ Hello 3 lần", "D. Lỗi do không dùng biến i"]),
    ("Câu 10: Vòng lặp `while` sẽ chạy khi nào?", ["A. Khi điều kiện bị Sai (False)", "B. Khi điều kiện còn Đúng (True)", "C. Chạy chính xác 10 lần", "D. Chỉ chạy khi có dấu [ ]"]),
    ("Câu 11: Đoạn code sau lặp mấy lần?\ni = 0\nwhile i < 3:\n    i = i + 1", ["A. 2 lần", "B. 4 lần", "C. Vô tận", "D. 3 lần"]),
    ("Câu 12: Đâu là nguyên nhân phổ biến gây ra 'vòng lặp vô tận' (infinite loop) với `while`?", ["A. Quên cập nhật biến điều kiện (ví dụ quên lệnh i = i + 1)", "B. Sử dụng quá nhiều List", "C. Máy tính bị virus", "D. Do dùng Python 3 thay vì Python 2"]),
    ("Câu 13: Lệnh `break` có tác dụng gì trong vòng lặp?", ["A. Tạm dừng chương trình 1 giây", "B. Thoát ngay lập tức khỏi vòng lặp", "C. Xóa phần tử trong List", "D. Chạy lại vòng lặp từ đầu"]),
    ("Câu 14: Lệnh `continue` có tác dụng gì trong vòng lặp?", ["A. Thoát vòng lặp giống break", "B. Tắt màn hình máy tính", "C. Bỏ qua phần còn lại của vòng lặp hiện tại và nhảy sang vòng kế tiếp", "D. Xóa dữ liệu ổ cứng"]),
    ("Câu 15: Có đoạn code:\ncolors = ['Đỏ', 'Xanh']\nfor c in colors:\n    print(c)\nKết quả sẽ in ra?", ["A. Đỏ Xanh", "B. Đỏ (xuống dòng) Xanh", "C. ['Đỏ', 'Xanh']", "D. Lỗi cú pháp"]),
    ("Câu 16: Chỉ số `[-1]` trong `inventory[-1]` mang ý nghĩa gì?", ["A. Lấy phần tử bị lỗi", "B. Xóa phần tử đầu tiên", "C. Lấy phần tử cuối cùng của List", "D. Khởi động lại vòng lặp"]),
    ("Câu 17: Cú pháp tạo một vòng lặp `while` vô tận có chủ ý là gì?", ["A. while True:", "B. while False:", "C. for True:", "D. if True:"]),
    ("Câu 18: Để đếm xem List `cars` có bao nhiêu xe, ta dùng hàm gì?", ["A. length(cars)", "B. len(cars)", "C. count(cars)", "D. size(cars)"]),
    ("Câu 19: Kết quả của `print([1, 2] + [3, 4])` là?", ["A. [4, 6]", "B. [1, 2, 3, 4]", "C. 10", "D. Lỗi"]),
    ("Câu 20: Dấu hai chấm (:) ở cuối lệnh `for` hoặc `while` dùng để làm gì?", ["A. Ký hiệu kết thúc câu lệnh", "B. Bắt đầu khối lệnh lùi lề (block)", "C. Để in ra dấu hai chấm", "D. Python không bắt buộc dấu hai chấm"])
]

for idx, (q, opts) in enumerate(questions):
    p = doc2.add_paragraph()
    r = p.add_run(q)
    r.font.name = 'Arial'
    r.font.color.rgb = BODY_COLOR
    r.bold = True
    for opt in opts:
        add_body_paragraph(doc2, opt)
    add_body_paragraph(doc2, '')

doc2.add_page_break()
h_ans = doc2.add_heading('ĐÁP ÁN', level=1)
apply_heading_style(h_ans, 1)

answers = [
    "Câu 1: B", "Câu 2: A", "Câu 3: C (Cam vì đếm từ 0)", "Câu 4: B", 
    "Câu 5: B", "Câu 6: B", "Câu 7: A", "Câu 8: B", "Câu 9: B", "Câu 10: B",
    "Câu 11: D", "Câu 12: A", "Câu 13: B", "Câu 14: C", "Câu 15: B", "Câu 16: C",
    "Câu 17: A", "Câu 18: B", "Câu 19: B", "Câu 20: B"
]

for ans in answers:
    add_body_paragraph(doc2, ans)

doc2.save(os.path.join(dir_path, 'Buoi_2_Trac_Nghiem_Python.docx'))
print("Session 2 Files created successfully with Forest Canopy theme.")
