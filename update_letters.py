import docx

# Letter 1 content
l1_text = """LETTER OF RECOMMENDATION
June 23, 2026
Dear Admissions Committee, 
It is a pleasure to write this letter of recommendation for Luong Nguyen Minh as he applies for the Electrical Engineering Program at Ming Chuan University. As his instructor for several technical modules at Aptech, I have had chances to assess his practical engineering skills and academic potential.
Minh demonstrated a solid practical approach to problem-solving. In our lab sessions and project work, he showed capable coding skills and resolved system issues effectively. When he encountered software bugs, Minh was persistent. He has shown patience and determination suitable for the engineering field.
Additionally, he possesses good language skills that will assist him in an international academic setting: he is capable of using English for technical research. Throughout his studies, he did not just depend on textbooks. Instead, he researched international technical documents, developer forums, and English datasheets to tackle programming problems. His ability to understand technical English enables him to gain new knowledge independently.
In addition to his technical and language skills, Minh is a collaborative team player. He is open-minded and eager to explore the field of Electrical Engineering. His programming background will serve him well.
I believe that Minh's practical IT skills, language proficiency, and work ethic make him a solid candidate for your program. I recommend him to your university.
Sincerely,  
[Name]
[Title] 
Aptech Computer Education
[Work email]"""

# Letter 2 content
l2_text = """LETTER OF RECOMMENDATION
June 23, 2026
Dear Admissions Committee, 
I am pleased to recommend Luong Nguyen Minh for admission to the Electrical Engineering Program at Ming Chuan University.
I have had the opportunity of teaching and working with Minh during his studies at Aptech. Throughout this time, he showed capable analytical thinking, logical reasoning, and an interest in technical subjects.
He was a diligent learner who took his assignments seriously. He grasped concepts well and applied them in practical exercises and projects. His problem-solving skills and willingness to explore new knowledge are commendable.
Beyond his academic performance, he maintained a positive attitude toward learning. He was respectful, cooperative, and open to feedback.
Based on my observations, I believe he has the discipline and determination needed to succeed in your academic environment. His background in information technology will help him contribute to the Electrical Engineering department at Ming Chuan University.
Therefore, I recommend him for admission.
Sincerely,  
[Name]
[Title] 
Aptech Computer Education
[Work email]"""

doc1 = docx.Document()
for p in l1_text.split('\n'):
    doc1.add_paragraph(p)
doc1.save(r"c:\code\Cecomtech\recommend_letter_1.docx")

doc2 = docx.Document()
for p in l2_text.split('\n'):
    doc2.add_paragraph(p)
doc2.save(r"c:\code\Cecomtech\recommend_letter_2.docx")
print("Letters updated successfully.")
