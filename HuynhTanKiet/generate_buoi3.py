import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(20, 100, 20)  # Màu xanh lá đậm để phân biệt
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p.paragraph_format.left_indent = Pt(20)

# --- Tạo file Lý Thuyết ---
doc_lt = docx.Document()

title = doc_lt.add_heading('LẬP TRÌNH PYTHON: BUỔI 3 - HÀM, TỪ ĐIỂN, TUPLE, SET VÀ MODULE', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

p_muc_tieu = doc_lt.add_paragraph()
p_muc_tieu.add_run('Mục tiêu: ').bold = True
p_muc_tieu.add_run('Hiểu cách đóng gói code qua Hàm (Function). Làm chủ các cấu trúc dữ liệu Từ điển (Dictionary), Tuple và Set. Biết cách sử dụng thư viện (Module) có sẵn trong Python. (Thời lượng: 6 giờ)\n')

# Phần 1: Hàm
h1 = doc_lt.add_heading('Phần 1: Hàm (Functions) - Đóng gói và tái sử dụng', level=2)
for run in h1.runs: run.font.name = 'Arial'

doc_lt.add_paragraph('1.1 Tại sao chúng ta cần Hàm?')
doc_lt.add_paragraph('Hàm (Function) giúp bạn "đóng gói" một đoạn code lại, đặt tên và tái sử dụng nhiều lần mà không cần viết lại (Nguyên tắc DRY - Don\'t Repeat Yourself).')

doc_lt.add_paragraph('1.2 Cú pháp cơ bản và Tham số mặc định (Default arguments)')
add_code_block(doc_lt, "def greet_person(name, greeting='Xin chào'):\n    print(f'{greeting} {name}!')\n\ngreet_person('Minh')              # In ra: Xin chào Minh!\ngreet_person('An', 'Chào buổi sáng') # In ra: Chào buổi sáng An!")

doc_lt.add_paragraph('1.3 Tham số biến thiên (*args và **kwargs)')
doc_lt.add_paragraph('Khi bạn không biết trước có bao nhiêu tham số truyền vào hàm, sử dụng *args cho danh sách tham số, và **kwargs cho tham số dạng khóa-giá trị.')
add_code_block(doc_lt, "def calculate_sum(*args):\n    return sum(args)\n\nprint(calculate_sum(1, 2, 3, 4))  # In ra: 10")

doc_lt.add_paragraph('1.4 Hàm ẩn danh (Lambda)')
doc_lt.add_paragraph('Dùng để viết hàm ngắn gọn chỉ trên 1 dòng.')
add_code_block(doc_lt, "double_value = lambda x: x * 2\nprint(double_value(5))  # In ra: 10")

# Phần 2: Dictionary
h2 = doc_lt.add_heading('Phần 2: Từ điển (Dictionary) - Lưu trữ dữ liệu theo cặp', level=2)
for run in h2.runs: run.font.name = 'Arial'

doc_lt.add_paragraph('Dictionary lưu dữ liệu theo cấu trúc Khóa - Giá trị (Key - Value).')
add_code_block(doc_lt, "student = {\n    'name': 'Minh',\n    'age': 15,\n    'math_score': 9.5\n}\n# Thêm/Sửa\nstudent['literature_score'] = 8.0\n# Xóa\nstudent.pop('age')\n# Lặp qua Dictionary\nfor key, val in student.items():\n    print(f'{key}: {val}')")

# Phần 3: Tuple và Set
h3 = doc_lt.add_heading('Phần 3: Tuple và Set', level=2)
for run in h3.runs: run.font.name = 'Arial'

doc_lt.add_paragraph('3.1 Tuple - Danh sách không thể thay đổi (Immutable List)')
doc_lt.add_paragraph('Giống như List, nhưng sử dụng ngoặc tròn () và KHÔNG THỂ sửa đổi sau khi tạo.')
add_code_block(doc_lt, "coordinates = (10.0, 20.0)\nprint(coordinates[0])  # 10.0\n# coordinates[0] = 15.0  --> LỖI!")

doc_lt.add_paragraph('3.2 Set - Tập hợp toán học')
doc_lt.add_paragraph('Sử dụng ngoặc nhọn {}, chứa các phần tử DUY NHẤT (không trùng lặp) và KHÔNG CÓ THỨ TỰ.')
add_code_block(doc_lt, "unique_numbers = {1, 2, 2, 3, 4, 4}\nprint(unique_numbers)  # In ra: {1, 2, 3, 4}")

# Phần 4: Module
h4 = doc_lt.add_heading('Phần 4: Module (Sử dụng Thư viện)', level=2)
for run in h4.runs: run.font.name = 'Arial'

doc_lt.add_paragraph('Python có sẵn rất nhiều module (thư viện). Sử dụng từ khóa "import" để nạp module.')
add_code_block(doc_lt, "import math\nimport random\n\nprint(math.sqrt(16))          # In ra: 4.0\nprint(random.randint(1, 10))  # Số ngẫu nhiên từ 1 đến 10")

doc_lt.save('Buoi_3_Ly_Thuyet_Python.docx')

# --- Tạo file Trắc Nghiệm ---
doc_tn = docx.Document()

title_tn = doc_tn.add_heading('BÀI TẬP TRẮC NGHIỆM PYTHON - BUỔI 3 (50 CÂU)', level=1)
title_tn.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title_tn.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

doc_tn.add_paragraph('Hãy chọn đáp án đúng nhất cho các câu hỏi sau đây:\n')

cau_hoi = [
    # Functions Basics (1-10)
    {"q": "Câu 1: Từ khóa nào được sử dụng để định nghĩa một hàm trong Python?", "options": ["A. function", "B. def", "C. define", "D. func"]},
    {"q": "Câu 2: Đâu là cách gọi hàm chính xác nếu hàm có tên là `my_function`?", "options": ["A. call my_function()", "B. my_function", "C. my_function()", "D. run my_function"]},
    {"q": "Câu 3: Để trả về một giá trị từ một hàm, ta dùng từ khóa nào?", "options": ["A. yield", "B. return", "C. get", "D. output"]},
    {"q": "Câu 4: Nếu một hàm không có câu lệnh `return`, giá trị mặc định nào sẽ được trả về khi hàm kết thúc?", "options": ["A. 0", "B. Chuỗi rỗng \"\"", "C. None", "D. False"]},
    {"q": "Câu 5: Đầu vào của hàm được gọi là gì?", "options": ["A. Variables", "B. Loops", "C. Parameters / Arguments", "D. Methods"]},
    {"q": "Câu 6: Hàm trong Python có bắt buộc phải trả về một giá trị không?", "options": ["A. Có, luôn luôn", "B. Không bắt buộc", "C. Chỉ bắt buộc nếu hàm có tham số", "D. Chỉ bắt buộc nếu dùng print"]},
    {"q": "Câu 7: Dấu nào ngăn cách các tham số trong định nghĩa hàm?", "options": ["A. Dấu chấm phẩy (;)", "B. Dấu hai chấm (:)", "C. Dấu phẩy (,)", "D. Dấu chấm (.)"]},
    {"q": "Câu 8: Ở cuối dòng định nghĩa hàm (`def func():`), ký tự nào bắt buộc phải có?", "options": ["A. ;", "B. :", "C. {", "D. }"]},
    {"q": "Câu 9: Thân hàm (phần code bên trong hàm) phải tuân thủ quy tắc nào?", "options": ["A. Phải được bao trong ngoặc nhọn {}", "B. Phải được thụt lề (indentation)", "C. Phải kết thúc bằng dấu chấm phẩy", "D. Phải viết in hoa"]},
    {"q": "Câu 10: Tên hàm nào sau đây là KHÔNG hợp lệ?", "options": ["A. my_func()", "B. 1st_function()", "C. _secret_func()", "D. calculateSum()"]},
    # Functions Advanced (11-20)
    {"q": "Câu 11: Tham số mặc định (default parameter) được định nghĩa như thế nào trong hàm?", "options": ["A. def func(a default 10):", "B. def func(a=10):", "C. def func(a: 10):", "D. def func(a == 10):"]},
    {"q": "Câu 12: Khi gọi hàm `func(a=5, b=10)`, cách truyền đối số này gọi là gì?", "options": ["A. Default arguments", "B. Keyword arguments", "C. Positional arguments", "D. Required arguments"]},
    {"q": "Câu 13: `*args` trong định nghĩa hàm dùng để làm gì?", "options": ["A. Nhận một danh sách tham số vô hạn", "B. Yêu cầu một tham số kiểu số", "C. Nhân đôi giá trị tham số", "D. Định nghĩa biến cục bộ"]},
    {"q": "Câu 14: `**kwargs` được truyền vào hàm dưới dạng cấu trúc dữ liệu nào?", "options": ["A. List", "B. Tuple", "C. Dictionary", "D. Set"]},
    {"q": "Câu 15: Từ khóa nào dùng để tạo hàm ẩn danh (hàm trên một dòng)?", "options": ["A. def", "B. anon", "C. lambda", "D. inline"]},
    {"q": "Câu 16: Biến được khai báo bên trong hàm gọi là gì?", "options": ["A. Biến toàn cục (Global variable)", "B. Biến cục bộ (Local variable)", "C. Biến tĩnh (Static variable)", "D. Biến hằng (Constant)"]},
    {"q": "Câu 17: Để sử dụng và sửa đổi một biến toàn cục bên trong hàm, ta cần dùng từ khóa nào?", "options": ["A. global", "B. local", "C. static", "D. self"]},
    {"q": "Câu 18: Kết quả của đoạn code sau là gì?", "code": "def multiply(a, b=2):\n    return a * b\nprint(multiply(5))", "options": ["A. 5", "B. 2", "C. 10", "D. Báo lỗi do thiếu tham số"]},
    {"q": "Câu 19: Trong Python, ta có thể viết một hàm bên trong một hàm khác không?", "options": ["A. Có", "B. Không"]},
    {"q": "Câu 20: Câu lệnh `pass` trong hàm dùng để làm gì?", "options": ["A. Bỏ qua một lần lặp", "B. Dừng hàm lại", "C. Giữ chỗ (placeholder) khi hàm chưa có nội dung để tránh lỗi cú pháp", "D. Trả về giá trị pass"]},
    # Dictionary (21-30)
    {"q": "Câu 21: Dictionary (Từ điển) trong Python được khai báo bằng cặp dấu ngoặc nào?", "options": ["A. ( )", "B. [ ]", "C. { }", "D. < >"]},
    {"q": "Câu 22: Cú pháp khởi tạo một Dictionary rỗng?", "options": ["A. d = ()", "B. d = []", "C. d = {}", "D. d = empty"]},
    {"q": "Câu 23: Trong Dictionary, các phần tử được lưu trữ theo cấu trúc nào?", "options": ["A. Chỉ số - Giá trị", "B. Khóa - Giá trị (Key-Value)", "C. Hàng - Cột", "D. Danh sách liên kết"]},
    {"q": "Câu 24: Khóa (Key) trong Dictionary phải tuân thủ điều kiện gì?", "options": ["A. Phải là một số nguyên", "B. Phải duy nhất (không trùng lặp)", "C. Phải là chuỗi", "D. Có thể trùng lặp"]},
    {"q": "Câu 25: Để lấy giá trị của khóa 'age' từ `info = {'name': 'An', 'age': 20}`, ta dùng:", "options": ["A. info[1]", "B. info('age')", "C. info.age", "D. info['age']"]},
    {"q": "Câu 26: Phương thức nào dùng để lấy ra tất cả các Khóa trong Dictionary?", "options": ["A. .values()", "B. .keys()", "C. .items()", "D. .get()"]},
    {"q": "Câu 27: Phương thức nào dùng để lấy ra tất cả các Giá trị trong Dictionary?", "options": ["A. .values()", "B. .keys()", "C. .items()", "D. .get()"]},
    {"q": "Câu 28: Phương thức nào lấy ra các cặp (Khóa, Giá trị) dưới dạng danh sách Tuple?", "options": ["A. .values()", "B. .keys()", "C. .items()", "D. .all()"]},
    {"q": "Câu 29: Toán tử nào kiểm tra một khóa có tồn tại trong Dictionary hay không?", "options": ["A. exists", "B. in", "C. has", "D. contains"]},
    {"q": "Câu 30: Phương thức nào dùng để xóa và trả về giá trị của một khóa trong Dictionary?", "options": ["A. delete()", "B. remove()", "C. pop()", "D. discard()"]},
    # Tuple & Set (31-40)
    {"q": "Câu 31: Tuple trong Python được khai báo bằng dấu ngoặc nào?", "options": ["A. ( )", "B. [ ]", "C. { }", "D. < >"]},
    {"q": "Câu 32: Điểm khác biệt lớn nhất giữa List và Tuple là gì?", "options": ["A. Tuple không chứa được chuỗi", "B. Tuple không thể thay đổi giá trị (Immutable) sau khi tạo", "C. List nhanh hơn Tuple", "D. Tuple không dùng index"]},
    {"q": "Câu 33: Cách tạo một Tuple chỉ có MỘT phần tử?", "options": ["A. t = (1)", "B. t = [1]", "C. t = (1,)", "D. t = {1}"]},
    {"q": "Câu 34: Set (Tập hợp) trong Python được khai báo bằng dấu ngoặc nào?", "options": ["A. ( )", "B. [ ]", "C. { }", "D. < >"]},
    {"q": "Câu 35: Đặc điểm nổi bật nhất của Set là gì?", "options": ["A. Lưu trữ dữ liệu siêu lớn", "B. Các phần tử là duy nhất (không trùng lặp)", "C. Có thể truy cập bằng index", "D. Giống hệt List"]},
    {"q": "Câu 36: Lệnh nào dùng để thêm một phần tử vào Set `s`?", "options": ["A. s.append(5)", "B. s.insert(5)", "C. s.push(5)", "D. s.add(5)"]},
    {"q": "Câu 37: Lệnh nào chuyển đổi một List `[1, 2, 2, 3]` thành một Set để loại bỏ phần tử trùng?", "options": ["A. tuple([1, 2, 2, 3])", "B. set([1, 2, 2, 3])", "C. dict([1, 2, 2, 3])", "D. unique([1, 2, 2, 3])"]},
    {"q": "Câu 38: Kết quả của đoạn code: `print(len({1, 1, 2, 2, 3}))` là gì?", "options": ["A. 5", "B. 3", "C. 4", "D. 1"]},
    {"q": "Câu 39: Set có hỗ trợ truy cập phần tử qua chỉ số (index), ví dụ `my_set[0]` không?", "options": ["A. Có", "B. Không"]},
    {"q": "Câu 40: Tuple có các phương thức như `.append()` hay `.remove()` không?", "options": ["A. Có", "B. Không"]},
    # Modules & General (41-50)
    {"q": "Câu 41: Từ khóa nào dùng để nạp một Module (Thư viện) vào chương trình Python?", "options": ["A. include", "B. require", "C. import", "D. using"]},
    {"q": "Câu 42: Để chỉ nạp hàm `sqrt` từ thư viện `math`, ta dùng cú pháp nào?", "options": ["A. import sqrt from math", "B. from math import sqrt", "C. include math.sqrt", "D. math.sqrt()"]},
    {"q": "Câu 43: Thư viện nào trong Python có sẵn (built-in) cung cấp các hàm tạo số ngẫu nhiên?", "options": ["A. math", "B. time", "C. random", "D. os"]},
    {"q": "Câu 44: Hàm `random.randint(1, 5)` sẽ trả về các số ngẫu nhiên nằm trong khoảng nào?", "options": ["A. 1, 2, 3, 4", "B. 1, 2, 3, 4, 5", "C. 0, 1, 2, 3, 4", "D. 2, 3, 4"]},
    {"q": "Câu 45: Nếu muốn gán một bí danh (alias) ngắn hơn cho thư viện khi import, ta làm thế nào?", "options": ["A. import math as m", "B. import math = m", "C. from math alias m", "D. import m from math"]},
    {"q": "Câu 46: Trong Dictionary, lệnh `student.update({'score': 10})` dùng để làm gì?", "options": ["A. Xóa khóa score", "B. Thêm hoặc cập nhật khóa score thành 10", "C. Lấy ra khóa score", "D. Đổi tên khóa thành score"]},
    {"q": "Câu 47: Để xóa toàn bộ phần tử trong một List, Dictionary hoặc Set, phương thức nào được dùng chung?", "options": ["A. .remove_all()", "B. .delete()", "C. .clear()", "D. .empty()"]},
    {"q": "Câu 48: Tuple có thể làm Khóa (Key) trong Dictionary không?", "options": ["A. Có, vì Tuple không thay đổi (immutable)", "B. Không, chỉ có chuỗi mới làm Khóa được"]},
    {"q": "Câu 49: Kết quả của đoạn code sau là gì?", "code": "def func(x=[]):\n    x.append(1)\n    return x\nprint(func())\nprint(func())", "options": ["A. [1] và [1]", "B. [1] và [1, 1]", "C. Lỗi cú pháp", "D. [1] và []"]},
    {"q": "Câu 50: Đâu là cấu trúc dữ liệu thích hợp nhất để lưu thông tin về tọa độ (x, y) trên bản đồ mà không bị thay đổi?", "options": ["A. List", "B. Dictionary", "C. Set", "D. Tuple"]}
]

for ch in cau_hoi:
    p = doc_tn.add_paragraph(ch['q'])
    p.runs[0].bold = True
    
    if "code" in ch:
        add_code_block(doc_tn, ch["code"])
        
    for opt in ch['options']:
        doc_tn.add_paragraph(opt)
    doc_tn.add_paragraph() # Dòng trống

doc_tn.save('Buoi_3_Trac_Nghiem_Python.docx')
