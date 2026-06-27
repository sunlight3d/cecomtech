import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

dir_path = r"c:\code\Cecomtech\HuynhTanKiet"
os.makedirs(dir_path, exist_ok=True)

# Tech Innovation Theme
HEADING_COLOR = RGBColor(0x00, 0x50, 0x9E) # Modern Tech Blue
BODY_COLOR = RGBColor(0x2C, 0x3E, 0x50)    # Slate Dark Blue/Gray
CODE_COLOR = RGBColor(0x00, 0x80, 0x80)    # Teal

def apply_heading_style(heading, level):
    for run in heading.runs:
        run.font.name = 'Trebuchet MS'
        run.font.color.rgb = HEADING_COLOR
        if level == 0:
            run.font.size = Pt(22)
            run.bold = True
        elif level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True

def add_body_paragraph(doc, text, bold_prefix=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Arial'
        r.font.color.rgb = BODY_COLOR
        
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.color.rgb = BODY_COLOR
    else:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.color.rgb = BODY_COLOR
    return p

def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.color.rgb = CODE_COLOR
    r.bold = True
    p.paragraph_format.left_indent = Pt(20)
    return p

# --- DOCUMENT 1: THEORY ---
doc1 = Document()
h0 = doc1.add_heading('LẬP TRÌNH PYTHON: BUỔI 2 - DANH SÁCH & VÒNG LẶP', 0)
h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
apply_heading_style(h0, 0)

add_body_paragraph(doc1, 'Mục tiêu: Hiểu sâu sắc về cấu trúc dữ liệu Danh sách (List), sử dụng thành thạo Vòng lặp (For, While) để tự động hóa các tác vụ lặp đi lặp lại. Xây dựng nền tảng tư duy xử lý dữ liệu số lượng lớn. (Thời lượng: 4 giờ)\n')

h1 = doc1.add_heading('Phần 1: Danh sách (List) - Chiếc Ba Lô Vạn Năng (90 phút)', level=1)
apply_heading_style(h1, 1)

add_body_paragraph(doc1, '1. Tại sao lại cần Danh sách?', bold_prefix='1.1 ')
add_body_paragraph(doc1, 'Hãy tưởng tượng bạn đang viết một phần mềm quản lý điểm thi cho 1000 học sinh. Nếu chỉ dùng "Biến" (như buổi 1), bạn sẽ phải tạo ra 1000 biến khác nhau (score_1 = 8, score_2 = 9,... score_1000 = 7). Việc này quá tốn thời gian và dễ sai sót. Danh sách (List) sinh ra để giải quyết vấn đề này. List giống như một toa tàu hỏa có rất nhiều khoang, mỗi khoang có thể chứa một dữ liệu khác nhau.')

try:
    doc1.add_picture(r"C:\Users\ADMIN\.gemini\antigravity\brain\6c9a90d7-9cf7-4601-aaaa-593d205ae5d0\python_list_concept_1782443502927.png", width=Inches(5))
except Exception as e:
    print(f"Error loading image 1: {e}")

add_body_paragraph(doc1, '1.2 Cách tạo ra một Danh sách')
add_body_paragraph(doc1, 'Để tạo một danh sách, chúng ta sử dụng cặp dấu ngoặc vuông [ ]. Các phần tử bên trong được phân cách bằng dấu phẩy.')
add_code_paragraph(doc1, "students = ['Minh', 'Khoa', 'Lan', 'Huong']\nscores = [8.5, 9.0, 7.5, 10.0]\nmixed_data = ['Yasuo', 15, True, 3.14]")
add_body_paragraph(doc1, 'Bạn thấy đấy, một List có thể chứa chữ (chuỗi), số, hoặc thậm chí là Boolean (True/False) cùng một lúc. Nó rất linh hoạt.')

add_body_paragraph(doc1, '1.3 Chỉ số (Index) - Địa chỉ của từng khoang tàu')
add_body_paragraph(doc1, 'Máy tính đếm từ số 0 chứ không phải số 1. Phần tử đầu tiên trong List nằm ở vị trí số 0.')
add_code_paragraph(doc1, "fruits = ['Apple', 'Orange', 'Mango', 'Grape']\nprint(fruits[0])  # Sẽ in ra: Apple\nprint(fruits[2])  # Sẽ in ra: Mango")
add_body_paragraph(doc1, 'Nếu bạn dùng số âm, Python sẽ đếm ngược từ cuối danh sách:')
add_code_paragraph(doc1, "print(fruits[-1])  # Sẽ in ra: Grape (phần tử cuối cùng)\nprint(fruits[-2])  # Sẽ in ra: Mango")

add_body_paragraph(doc1, '1.4 Thêm đồ vật vào Danh sách')
add_body_paragraph(doc1, 'Khi có học sinh mới chuyển đến, ta cần thêm vào danh sách. Có 2 cách chính:')
add_body_paragraph(doc1, '- Hàm append(x):', bold_prefix='Thêm vào cuối (append): ')
add_code_paragraph(doc1, "fruits.append('Watermelon')\nprint(fruits)  # ['Apple', 'Orange', 'Mango', 'Grape', 'Watermelon']")
add_body_paragraph(doc1, '- Hàm insert(index, x):', bold_prefix='Chèn vào vị trí bất kỳ (insert): ')
add_code_paragraph(doc1, "fruits.insert(1, 'Durian')  # Chèn Durian vào vị trí số 1")

add_body_paragraph(doc1, '1.5 Xóa đồ vật khỏi Danh sách')
add_body_paragraph(doc1, '- Hàm remove(giá_trị):', bold_prefix='Xóa theo tên: ')
add_code_paragraph(doc1, "fruits.remove('Orange')  # Tìm chữ Orange và xóa nó")
add_body_paragraph(doc1, '- Hàm pop(index):', bold_prefix='Xóa theo vị trí (Index): ')
add_code_paragraph(doc1, "fruits.pop(0)  # Xóa phần tử ở vị trí 0 (Apple)")

add_body_paragraph(doc1, '1.6 Cắt danh sách (Slicing) - Tính năng cực kỳ mạnh mẽ của Python')
add_body_paragraph(doc1, 'Giả sử bạn có 100 học sinh, bạn chỉ muốn lấy ra 3 bạn đầu tiên. Bạn dùng dấu hai chấm [start:end].')
add_code_paragraph(doc1, "data_list = ['A', 'B', 'C', 'D', 'E']\nprint(data_list[0:3])  # Lấy từ vị trí 0 đến trước vị trí 3. Kết quả: ['A', 'B', 'C']\nprint(data_list[2:])   # Lấy từ vị trí 2 đến hết. Kết quả: ['C', 'D', 'E']")

h1 = doc1.add_heading('Phần 2: Vòng lặp FOR - Cỗ máy tự động hóa (90 phút)', level=1)
apply_heading_style(h1, 1)

try:
    doc1.add_picture(r"C:\Users\ADMIN\.gemini\antigravity\brain\6c9a90d7-9cf7-4601-aaaa-593d205ae5d0\python_loop_concept_1782443514848.png", width=Inches(5))
except Exception as e:
    print(f"Error loading image 2: {e}")

add_body_paragraph(doc1, '2.1 Khái niệm về Vòng lặp')
add_body_paragraph(doc1, 'Nếu tôi yêu cầu bạn in chữ "Hello" 10 lần. Bạn có thể copy lệnh print 10 lần. Nhưng nếu tôi yêu cầu 10,000 lần thì sao? Vòng lặp ra đời để thay con người làm những việc nhàm chán, lặp đi lặp lại một cách cực kỳ nhanh chóng.')

add_body_paragraph(doc1, '2.2 Vòng lặp For duyệt qua Danh sách')
add_body_paragraph(doc1, 'Đây là sức mạnh lớn nhất của vòng lặp For. Nó có thể đi qua từng phần tử trong một List để xử lý.')
add_code_paragraph(doc1, "cart = ['Mouse', 'Keyboard', 'Headset']\nfor item in cart:\n    print('Packing item:', item)")
add_body_paragraph(doc1, 'Máy tính sẽ làm như sau:\n- Lần 1: item = "Mouse" -> in ra "Packing item: Mouse"\n- Lần 2: item = "Keyboard" -> in ra "Packing item: Keyboard"\n- Lần 3: item = "Headset" -> in ra "Packing item: Headset"\nKhi hết danh sách, nó tự động dừng.')

add_body_paragraph(doc1, '2.3 Hàm range() - Máy tạo số tự động')
add_body_paragraph(doc1, 'Nếu bạn muốn vòng lặp chạy một số lần nhất định (ví dụ 5 lần) mà không có sẵn một danh sách, hãy dùng range(n). Nó sẽ tạo ra các số từ 0 đến n-1.')
add_code_paragraph(doc1, "for i in range(5):\n    print('Current count is:', i)")
add_body_paragraph(doc1, 'Bạn cũng có thể chỉ định điểm bắt đầu và kết thúc: range(1, 11) sẽ tạo các số từ 1 đến 10.')

add_body_paragraph(doc1, '2.4 Bài toán kinh điển: Tính tổng')
add_body_paragraph(doc1, 'Hãy tính tổng các con số trong một danh sách.')
add_code_paragraph(doc1, "scores = [7, 8, 9, 10]\ntotal_score = 0\nfor score in scores:\n    total_score = total_score + score\nprint('Total score is:', total_score)  # Kết quả là 34")

h1 = doc1.add_heading('Phần 3: Vòng lặp WHILE - Làm cho đến khi dừng (45 phút)', level=1)
apply_heading_style(h1, 1)

add_body_paragraph(doc1, '3.1 Sự khác biệt giữa For và While')
add_body_paragraph(doc1, '- For: Lặp khi biết trước số lần lặp (ví dụ lặp 10 lần, hoặc lặp hết các món trong danh sách).\n- While: Lặp dựa trên một điều kiện. Chừng nào điều kiện còn Đúng (True), nó cứ lặp mãi. Ta không cần biết trước số lần lặp.')

add_body_paragraph(doc1, '3.2 Ví dụ cơ bản')
add_code_paragraph(doc1, "energy = 0\nwhile energy < 100:\n    print('Charging...', energy)\n    energy = energy + 20\nprint('Battery full!')")

add_body_paragraph(doc1, '3.3 Cảnh báo: Vòng lặp vô tận (Infinite Loop)')
add_body_paragraph(doc1, 'Nếu bạn quên tăng biến energy ở ví dụ trên, energy sẽ luôn là 0. Điều kiện 0 < 100 luôn luôn ĐÚNG. Máy tính sẽ in ra chữ "Charging... 0" hàng triệu lần cho đến khi treo máy. Đây là lỗi phổ biến nhất của người mới học.')

add_body_paragraph(doc1, '3.4 Điều khiển luồng: break và continue')
add_body_paragraph(doc1, '- Lệnh break: Đập vỡ vòng lặp, thoát ra ngay lập tức.')
add_code_paragraph(doc1, "while True:  # Cố tình tạo vòng lặp vô tận\n    password = input('Enter password: ')\n    if password == '123':\n        print('Door opened!')\n        break  # Thoát ngay\n    else:\n        print('Wrong, try again!')")
add_body_paragraph(doc1, '- Lệnh continue: Bỏ qua các lệnh bên dưới nó, quay lại ngay từ đầu vòng lặp để lặp vòng tiếp theo.')
add_code_paragraph(doc1, "for number in range(1, 6):\n    if number == 3:\n        continue  # Bỏ qua số 3\n    print(number)  # In ra: 1, 2, 4, 5")

h1 = doc1.add_heading('Phần 4: Ứng dụng Thực tế - Dự án Quản lý (15 phút)', level=1)
apply_heading_style(h1, 1)
add_body_paragraph(doc1, 'Kết hợp tất cả những gì đã học, học sinh sẽ lập trình một phần mềm nhỏ.')
add_body_paragraph(doc1, 'Yêu cầu:')
add_body_paragraph(doc1, '1. Khởi tạo một danh sách inventory = [] (Rỗng).')
add_body_paragraph(doc1, '2. Dùng vòng lặp while True để tạo một menu chương trình quản lý.')
add_body_paragraph(doc1, '3. Nhập 1 để thêm mặt hàng vào kho (dùng append).')
add_body_paragraph(doc1, '4. Nhập 2 để hiển thị tất cả mặt hàng (dùng for).')
add_body_paragraph(doc1, '5. Nhập 3 để thoát (dùng break).')

file_theory = os.path.join(dir_path, 'Buoi_2_Ly_Thuyet_Python.docx')
try:
    doc1.save(file_theory)
except PermissionError:
    file_theory = os.path.join(dir_path, 'Buoi_2_Ly_Thuyet_Python_V3.docx')
    doc1.save(file_theory)

# --- DOCUMENT 2: MULTIPLE CHOICE ---
doc2 = Document()
h0 = doc2.add_heading('BÀI TẬP TRẮC NGHIỆM PYTHON - BUỔI 2', 0)
h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
apply_heading_style(h0, 0)

add_body_paragraph(doc2, 'Hãy chọn đáp án đúng nhất cho các câu hỏi sau đây:')

questions = [
    ("Câu 1: List (Danh sách) trong Python được khai báo bằng cặp dấu ngoặc nào?", ["A. ( )", "B. [ ]", "C. { }", "D. < >"]),
    ("Câu 2: Khai báo List nào sau đây là hợp lệ?", ["A. numbers = [1, 2, 3]", "B. numbers = (1, 2, 3)", "C. numbers = 1, 2, 3", "D. numbers = {1, 2, 3}"]),
    ("Câu 3: Trong List `fruits = ['Apple', 'Orange', 'Grape']`, lệnh `print(fruits[1])` sẽ in ra gì?", ["A. Apple", "B. Grape", "C. Orange", "D. Lỗi"]),
    ("Câu 4: Chỉ số (Index) của phần tử đầu tiên trong List là số mấy?", ["A. 1", "B. 0", "C. -1", "D. 2"]),
    ("Câu 5: Làm thế nào để thêm 'Banana' vào cuối List `fruits`?", ["A. fruits.add('Banana')", "B. fruits.append('Banana')", "C. fruits.insert('Banana')", "D. fruits.push('Banana')"]),
    ("Câu 6: Lệnh `remove('Apple')` dùng để làm gì?", ["A. Xóa tất cả các chữ Apple trong máy tính", "B. Xóa 'Apple' khỏi List", "C. Tìm vị trí của chữ 'Apple'", "D. Thêm chữ 'Apple' vào List"]),
    ("Câu 7: Vòng lặp `for` thường được sử dụng khi nào?", ["A. Khi biết trước số lần lặp hoặc duyệt qua một danh sách.", "B. Khi không biết trước số lần lặp.", "C. Khi muốn máy tính bị treo.", "D. Khi muốn tính tổng."]),
    ("Câu 8: Hàm `range(5)` sẽ tạo ra các số nào?", ["A. 1, 2, 3, 4, 5", "B. 0, 1, 2, 3, 4", "C. 0, 1, 2, 3, 4, 5", "D. 1, 2, 3, 4"]),
    ("Câu 9: Kết quả của đoạn code sau là gì?\nfor i in range(2):\n    print('Hello')", ["A. In chữ Hello 1 lần", "B. In chữ Hello 2 lần", "C. In chữ Hello 3 lần", "D. Lỗi do không dùng biến i"]),
    ("Câu 10: Vòng lặp `while` sẽ chạy khi nào?", ["A. Khi điều kiện bị Sai (False)", "B. Khi điều kiện còn Đúng (True)", "C. Chạy chính xác 10 lần", "D. Chỉ chạy khi có dấu [ ]"]),
    ("Câu 11: Đoạn code sau lặp mấy lần?\ni = 0\nwhile i < 3:\n    i = i + 1", ["A. 2 lần", "B. 4 lần", "C. Vô tận", "D. 3 lần"]),
    ("Câu 12: Đâu là nguyên nhân phổ biến gây ra 'vòng lặp vô tận' (infinite loop) với `while`?", ["A. Quên cập nhật biến điều kiện (ví dụ quên lệnh i = i + 1)", "B. Sử dụng quá nhiều List", "C. Máy tính bị virus", "D. Do dùng Python 3 thay vì Python 2"]),
    ("Câu 13: Lệnh `break` có tác dụng gì trong vòng lặp?", ["A. Tạm dừng chương trình 1 giây", "B. Thoát ngay lập tức khỏi vòng lặp", "C. Xóa phần tử trong List", "D. Chạy lại vòng lặp từ đầu"]),
    ("Câu 14: Lệnh `continue` có tác dụng gì trong vòng lặp?", ["A. Thoát vòng lặp giống break", "B. Tắt màn hình máy tính", "C. Bỏ qua phần còn lại của vòng lặp hiện tại và nhảy sang vòng kế tiếp", "D. Xóa dữ liệu ổ cứng"]),
    ("Câu 15: Có đoạn code:\ncolors = ['Red', 'Blue']\nfor c in colors:\n    print(c)\nKết quả sẽ in ra?", ["A. Red Blue", "B. Red (xuống dòng) Blue", "C. ['Red', 'Blue']", "D. Lỗi cú pháp"]),
    ("Câu 16: Chỉ số `[-1]` trong `inventory[-1]` mang ý nghĩa gì?", ["A. Lấy phần tử bị lỗi", "B. Xóa phần tử đầu tiên", "C. Lấy phần tử cuối cùng của List", "D. Khởi động lại vòng lặp"]),
    ("Câu 17: Cú pháp tạo một vòng lặp `while` vô tận có chủ ý là gì?", ["A. while True:", "B. while False:", "C. for True:", "D. if True:"]),
    ("Câu 18: Để đếm xem List `cars` có bao nhiêu xe, ta dùng hàm gì?", ["A. length(cars)", "B. len(cars)", "C. count(cars)", "D. size(cars)"]),
    ("Câu 19: Kết quả của `print([1, 2] + [3, 4])` là?", ["A. [4, 6]", "B. [1, 2, 3, 4]", "C. 10", "D. Lỗi"]),
    ("Câu 20: Dấu hai chấm (:) ở cuối lệnh `for` hoặc `while` dùng để làm gì?", ["A. Ký hiệu kết thúc câu lệnh", "B. Bắt đầu khối lệnh lùi lề (block)", "C. Để in ra dấu hai chấm", "D. Python không bắt buộc dấu hai chấm"])
]

for idx, (q, opts) in enumerate(questions):
    p = doc2.add_paragraph()
    r = p.add_run(q)
    r.font.name = 'Arial'
    r.font.color.rgb = BODY_COLOR
    r.bold = True
    for opt in opts:
        add_body_paragraph(doc2, opt)
    add_body_paragraph(doc2, '')

doc2.add_page_break()
h_ans = doc2.add_heading('ĐÁP ÁN', level=1)
apply_heading_style(h_ans, 1)

answers = [
    "Câu 1: B", "Câu 2: A", "Câu 3: C (Orange vì đếm từ 0)", "Câu 4: B", 
    "Câu 5: B", "Câu 6: B", "Câu 7: A", "Câu 8: B", "Câu 9: B", "Câu 10: B",
    "Câu 11: D", "Câu 12: A", "Câu 13: B", "Câu 14: C", "Câu 15: B", "Câu 16: C",
    "Câu 17: A", "Câu 18: B", "Câu 19: B", "Câu 20: B"
]

for ans in answers:
    add_body_paragraph(doc2, ans)

file_question = os.path.join(dir_path, 'Buoi_2_Trac_Nghiem_Python.docx')
try:
    doc2.save(file_question)
except PermissionError:
    file_question = os.path.join(dir_path, 'Buoi_2_Trac_Nghiem_Python_V3.docx')
    doc2.save(file_question)

print(f"Session 2 Files updated with English variables and images in {dir_path}.")
